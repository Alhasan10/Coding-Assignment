from docx import Document
from collections import Counter
import heapq
from math import log2


# Read the full story from a .docx file, including paragraphs and tables
def read_story(filename):
    doc = Document(filename)
    content = []

    # Extract paragraphs
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            content.append(paragraph.text.strip())

    # Extract tables
    for table_index, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    content.append(cell_text)

    # Combine all text into a single string
    full_text = "\n".join(content)
    return full_text


# Preprocess text
def preprocess_text(text):
    text = text.lower()
    text = text.replace("\n", "")  # Remove newline characters
    return text


# Calculate character frequencies
def calculate_frequencies(text):
    return Counter(text)


# Calculate probabilities
def calculate_probabilities(frequencies, total_chars):
    return {char: freq / total_chars for char, freq in frequencies.items()}


# Calculate entropy
def calculate_entropy(probabilities):
    return -sum(p * log2(p) for p in probabilities.values() if p > 0)


# Build Huffman tree and generate codes
def build_huffman_tree(frequencies):
    heap = [[weight, [char, ""]] for char, weight in frequencies.items()]
    heapq.heapify(heap)
    while len(heap) > 1:
        lo = heapq.heappop(heap)
        hi = heapq.heappop(heap)
        for pair in lo[1:]:
            pair[1] = '0' + pair[1]
        for pair in hi[1:]:
            pair[1] = '1' + pair[1]
        heapq.heappush(heap, [lo[0] + hi[0]] + lo[1:] + hi[1:])
    return sorted(heapq.heappop(heap)[1:], key=lambda p: (len(p[-1]), p))


# Calculate bits for ASCII and Huffman
def calculate_bits(frequencies, huffman_codes):
    total_ascii_bits = sum(freq * 8 for freq in frequencies.values())
    total_huffman_bits = sum(frequencies[char] * len(code) for char, code in huffman_codes)
    return total_ascii_bits, total_huffman_bits


# Main function
def main():
    file_path = r"C:\\Users\\hp\Downloads\\BZU\\1st Sem 4th\\Coding\\To+Build+A+Fire+by+Jack+London.docx"
    text = read_story(file_path)

    text = preprocess_text(text)

    frequencies = calculate_frequencies(text)
    total_chars = sum(frequencies.values())
    probabilities = calculate_probabilities(frequencies, total_chars)
    entropy = calculate_entropy(probabilities)

    huffman_codes = build_huffman_tree(frequencies)
    total_ascii_bits, total_huffman_bits = calculate_bits(frequencies, huffman_codes)
    compression_percentage = (total_ascii_bits - total_huffman_bits) / total_ascii_bits * 100

    # Calculate average bits per character
    avg_bits_per_char = total_huffman_bits / total_chars

    # Calculate sum of probabilities
    sum_probabilities = sum(probabilities.values())

    # Calculate entropy for alphabet characters only
    alphabet_probs = {char: prob for char, prob in probabilities.items() if char.isalpha()}
    alphabet_entropy = -sum(p * log2(p) for p in alphabet_probs.values() if p > 0)

    # Calculate informational efficiency
    informational_efficiency = (entropy / avg_bits_per_char) * 100

    # Sort characters alphabetically
    sorted_huffman_codes = sorted(huffman_codes, key=lambda x: x[0])

    # Print results in the desired format
    print(f"Sum of Probabilities = {sum_probabilities:.2f}")
    print(f"Entropy for all char = {entropy:.6f}")
    print(f"Entropy for alphabet char = {alphabet_entropy:.6f}")
    print(f"For ASCII coding, num of bits = {total_ascii_bits}")
    print(f"For Huffman coding, num of bits = {total_huffman_bits}")
    print(f"Average length of the code = {avg_bits_per_char:.6f}")
    print(f"Informational Efficiency = {informational_efficiency:.2f}%")
    print(f"Percentage of Compression = {compression_percentage:.2f}%\n")

    print(f"{'Symbol':<10}{'Probability':<15}{'Length Of Codeword':<20}{'Codeword':<10}")
    for char, code in sorted_huffman_codes:
        prob = probabilities[char]
        print(f"{repr(char):<10}{prob:<15.3f}{len(code):<20}{code:<10}")


if __name__ == "__main__":
    main()
